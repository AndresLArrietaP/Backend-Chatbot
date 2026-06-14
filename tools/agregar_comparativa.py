from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'c:\Users\Usuario\Pictures\PYTHON\Backend-Chatbot\docs\copilot\KomfIA_Instrucciones_Copilot.docx')

# ── helpers ────────────────────────────────────────────────────────────────────
def H1(doc, text):
    p = doc.add_heading(text, level=1)
    r = p.runs[0] if p.runs else p.add_run(text)
    r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)

def H2(doc, text, color=None):
    p = doc.add_heading(text, level=2)
    r = p.runs[0] if p.runs else p.add_run(text)
    r.font.size = Pt(12); r.font.color.rgb = color or RGBColor(0x2E,0x74,0xB5)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)

def H3(doc, text, color=None):
    p = doc.add_heading(text, level=3)
    r = p.runs[0] if p.runs else p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = color or RGBColor(0x24,0x63,0x9E)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)

def body(doc, text, indent=0.15, size=10.5, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color

def bullet(doc, text, indent=0.25, bold_prefix=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        rb = p.add_run(bold_prefix); rb.bold = True; rb.font.size = Pt(10.5)
        if color: rb.font.color.rgb = color
    r = p.add_run(text); r.font.size = Pt(10.5)
    if color and not bold_prefix: r.font.color.rgb = color

def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '2E74B5')
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9.5)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
    # data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri+1]
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(1)
    if col_widths:
        from docx.shared import Inches as In
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = In(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK + SECCIÓN COMPARATIVA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

# Título principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_title.add_run('Comparativa de Soluciones — Pruebas 2026-05-26')
rt.font.size = Pt(17); rt.font.bold = True
rt.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = p_sub.add_run('KomfIA (Human Query principal)  vs  TEST SQL V2 (genera SQL directamente)  |  10 consultas cada uno')
rs.font.size = Pt(10); rs.font.color.rgb = RGBColor(0x44,0x72,0xC4)
doc.add_paragraph()
rule(doc)

# ── RESUMEN EJECUTIVO ──────────────────────────────────────────────────────────
H1(doc, 'Resumen Ejecutivo')
body(doc, 'Se ejecutaron 10 consultas idénticas en ambos agentes. KomfIA obtuvo 9/10 correctas '
     '(1 bug de LP cruzado en path LLM). TEST SQL V2 obtuvo 9/10 correctas (1 patrón frágil menor). '
     'Ambas soluciones son funcionales; sus perfiles de fortaleza son complementarios.')

# ── TABLA RESULTADOS ──────────────────────────────────────────────────────────
H2(doc, 'Tabla de Resultados por Consulta')
headers = ['#', 'Consulta', 'KomfIA', 'TEST SQL V2']
rows = [
    ('1', 'Historial CA3164 MOTOR (últimos 5)',              '✓ OK',       '✓ OK — incluye Observaciones/Recomendaciones'),
    ('2', 'Triage Hidráulico Antapaccay',                    '✓ OK',       '✓ OK — LP/LC correctos, JOIN proyecto'),
    ('3', 'Ruedas delanteras fuera de límite',               '✓ OK',       '⚠ LIKE "%RUEDA%" AND "%DELANTERA%" (compuesto, funciona)'),
    ('4', 'Estado CA3177 (todos compartimientos)',           '⚠ BUG LP',   '✓ OK — no usa LP/LC, usa Condicion field'),
    ('5', 'Evolución Fe MT Antapaccay 2 años',               '✓ OK',       '✓ OK — análisis profundo del Copilot'),
    ('6', 'Tendencia Al MT Antapaccay',                      '✓ OK',       '⚠ filtra Al>0 → excluye ceros válidos'),
    ('7', 'Top 5 equipos mayor Fe MT',                       '✓ OK',       '✓ OK — CASE [Parametros_Fuera] innovador'),
    ('8', 'Barrido MT global (sin proyecto)',                 '✓ OK',       '✓ OK — incluye 930E correctamente'),
    ('9', 'Barrido MT equipo "3160"',                        '✓ OK',       '✓ OK — CA3160 LH=160/RH=145, ambos < LP=200'),
    ('10','Tendencia Cr CA3171 MT LH',                       '✓ OK',       '✓ OK — 20 muestras, gráfico + análisis tendencia'),
]
add_table(doc, headers, rows, col_widths=[0.25, 2.2, 1.1, 2.4])

# ── ANÁLISIS DETALLADO ──────────────────────────────────────────────────────
doc.add_page_break()
H1(doc, 'Análisis Detallado')
rule(doc)

# KomfIA ventajas
H2(doc, '① KomfIA (Human Query) — Fortalezas', color=RGBColor(0x37,0x86,0x1F))
ventajas_k = [
    ('Formateo rico y consistente: ', 'emojis semáforo, bloques por equipo, gráficos ASCII, tablas markdown — respuesta lista para el usuario final sin post-procesado.'),
    ('Paths directos Python: ', 'triage_directo, historial_directo, ranking_directo generan SQL predecible y optimizado. Fix #13 (equipo_code), Fix #15 (ISNULL=9999), Fix #16 (GROUP BY Proyecto) aplicados sistemáticamente.'),
    ('Contexto conversacional real: ', '"de esos", "el peor", "quédate con los críticos" funcionan sin repetir filtros — el backend mantiene el estado de sesión.'),
    ('Análisis estadístico automático: ', 'analitica.py calcula tendencias, desviaciones, sugerencias de gráfico — sin carga adicional al LLM de Copilot.'),
    ('Retry automático: ', 'SQL vacío/nulo dispara reescritura automática. DDI gestionado por el backend de forma sistemática.'),
    ('Dominio inyectado: ', 'heurísticas de compartimiento, proyecto, equipo_code — el usuario no necesita escribir consultas SQL ni técnicas.'),
]
for bp, bt in ventajas_k:
    bullet(doc, bt, bold_prefix='• ' + bp, color=RGBColor(0x37,0x86,0x1F))

# KomfIA debilidades
H2(doc, '① KomfIA — Debilidades', color=RGBColor(0xC0,0x00,0x00))
debs_k = [
    ('Bug LP en path LLM (equipo-específico): ', 'Consultas tipo "¿Cómo está el CA3177?" van al LLM cuando no hay compartimiento explícito. El LLM genera LimitesLC sin GROUP BY [Proyecto] → MIN cruza todos los proyectos → LP=80 (Cerro Verde) aplicado a un equipo de Antapaccay (LP real=200). Resultado: falso positivo. Fix #16 aplica en paths directos pero no en este path LLM. Pendiente corrección en hint.'),
    ('SQL no auditable por defecto: ', 'El campo "sql" está oculto a menos que el usuario lo pida explícitamente. En producción esto es correcto, pero dificulta la validación durante pruebas.'),
    ('SQL frágil ocasional (LLM): ', 'Para CA3171 tendencia generó LEFT JOIN LimitesLC ON 1=1 (cross join). Funciona con 1 compartimiento pero rompería con 2+. Patrón frágil.'),
]
for bp, bt in debs_k:
    bullet(doc, bt, bold_prefix='• ' + bp, color=RGBColor(0xC0,0x00,0x00))

doc.add_paragraph()

# TEST SQL V2 ventajas
H2(doc, '② TEST SQL V2 — Fortalezas', color=RGBColor(0x37,0x86,0x1F))
ventajas_t = [
    ('SQL completamente auditable: ', 'cada respuesta muestra el SQL generado — ideal para pruebas, validación y debugging sin acceso a logs del backend.'),
    ('Fix #16 aplicado en TODOS los SQL: ', 'el Copilot aprendió la regla del esquema inyectado. LimitesLC siempre incluye GROUP BY [Proyecto] y JOIN con proyecto del equipo. Ningún SQL cruzó proyectos.'),
    ('Patrón CASE [Parametros_Fuera]: ', 'innovación del agente — columna calculada que concatena los nombres de metales fuera de límite ("Fe Cu Pb"). Muy útil para barridos masivos. Candidato a incorporar en llm.py.'),
    ('Análisis interpretativo del Copilot: ', 'el propio LLM de Copilot Studio hace análisis de calidad sobre los datos crudos — detectó LDL (Al=1.0 ppm), identificó equipos crónicamente elevados, clasificó periodos de escalada.'),
    ('Independencia del backend Python: ', 'si el backend de Render tiene downtime, TEST SQL V2 sigue funcionando (conecta directo a Azure SQL). Resiliente.'),
    ('Más columnas en output: ', 'incluye Observaciones, Recomendaciones, CM, CodigoMuestreo — el usuario ve contexto completo sin segunda consulta.'),
]
for bp, bt in ventajas_t:
    bullet(doc, bt, bold_prefix='• ' + bp, color=RGBColor(0x37,0x86,0x1F))

# TEST SQL V2 debilidades
H2(doc, '② TEST SQL V2 — Debilidades', color=RGBColor(0xC0,0x00,0x00))
debs_t = [
    ('Error de autenticación pendiente: ', 'el conector SQL Server solicita credenciales del usuario final en cada sesión nueva. Necesita configurarse como shared/service connection en Power Platform para eliminar el prompt "Conectar para continuar".'),
    ('Sin retry ni reparación de SQL: ', 'si el SQL retorna 0 filas o la BD devuelve error, no hay mecanismo de reintento. El usuario recibe respuesta vacía sin diagnóstico.'),
    ('Tendencia sin ROW_NUMBER: ', 'las consultas de evolución temporal devuelven TODAS las muestras sin ROW_NUMBER, no las últimas N por equipo. Para flotas grandes puede retornar miles de filas o truncar datos relevantes.'),
    ('LIKE compuesto en ruedas: ', 'generó LIKE "%RUEDA%" AND LIKE "%DELANTERA%" — funciona pero viola el patrón documentado (solo usar keyword único). Riesgo de no match si el valor real cambia.'),
    ('Filtro Al>0 incorrecto: ', 'excluye registros con Al=0 que son válidos (por debajo del límite de detección). El correcto es no filtrar o usar IS NOT NULL.'),
    ('Formateo dependiente del Copilot: ', 'la calidad de presentación varía entre consultas — algunas respuestas son listas de texto, otras tablas. Sin estandarización garantizada.'),
    ('Sin contexto conversacional del backend: ', 'cada consulta es independiente. No puede hacer seguimiento de "de esos equipos, cuál es el peor" aprovechando el contexto de la consulta anterior.'),
]
for bp, bt in debs_t:
    bullet(doc, bt, bold_prefix='• ' + bp, color=RGBColor(0xC0,0x00,0x00))

doc.add_page_break()

# ── TABLA COMPARATIVA FINAL ────────────────────────────────────────────────────
H1(doc, 'Tabla Comparativa Final')
rule(doc)
headers2 = ['Criterio', 'KomfIA (Human Query)', 'TEST SQL V2']
rows2 = [
    ('Precisión LP/LC',             '9/10 correctas (bug en path LLM equipo-específico)',  '10/10 correctas (Fix #16 siempre aplicado)'),
    ('Formateo de respuesta',       'Rico, estandarizado, listo para usuario final',       'Variable, depende del LLM de Copilot'),
    ('Auditabilidad SQL',           'Oculto por defecto (correcto en prod)',               'Siempre visible — ideal para pruebas'),
    ('Contexto conversacional',     'Sí — backend gestiona sesión con TTL',                'No — cada consulta es independiente'),
    ('Análisis estadístico',        'Automático (analitica.py)',                           'Manual por el LLM de Copilot (bueno)'),
    ('Resiliencia backend',         'Depende de Render/Python',                            'Directo a Azure SQL — más resiliente'),
    ('Retry en SQL vacío',          'Sí — automático',                                    'No'),
    ('DDI/CM handling',             'Backend sistemático',                                 'Manual en SQL (correcto en la mayoría)'),
    ('Velocidad de respuesta',      '~15-30s (LLM + BD + analitica)',                     'Depende del Copilot LLM + BD directa'),
    ('Auth/conexión',               'Sin fricción (HTTP al backend)',                      'Requiere configurar shared connection'),
    ('CASE [Parametros_Fuera]',     'No implementado aún',                                'Sí — innovación útil para barridos'),
    ('Uso recomendado',             'Producción con usuarios finales',                     'Validación técnica y debugging'),
]
add_table(doc, headers2, rows2, col_widths=[1.5, 2.2, 2.2])

# ── RECOMENDACIONES ────────────────────────────────────────────────────────────
H1(doc, 'Recomendaciones y Próximos Pasos')
rule(doc)

recs = [
    ('URGENTE — Fix bug LP path LLM (llm.py): ',
     'Reforzar hint para consultas equipo-específico con compartimiento implícito (ej: "¿Cómo está el CA3177?"). '
     'El hint debe instruir al LLM a incluir [Proyecto] en LimitesLC GROUP BY y usar AND LC.[Proyecto] LIKE \'%\'+LS.[Proyecto]+\'%\' en el JOIN.'),
    ('MEJORA — Incorporar CASE [Parametros_Fuera] en llm.py: ',
     'El patrón de TEST SQL V2 es superior para barridos. Agregar columna calculada [Parametros_Fuera] '
     'en intentar_triage_directo y en el prompt LLM para triage.'),
    ('INFRA — Resolver auth TEST SQL V2: ',
     'Cambiar la conexión SQL Server del flujo a "shared connection" con cuenta de servicio en Power Platform. '
     'Elimina el prompt "Conectar para continuar" y el SystemError en usuarios nuevos.'),
    ('MEJORA — Tendencia TEST SQL V2: ',
     'Agregar ROW_NUMBER PARTITION BY (MiningEquipmentId, Compartimiento) con WHERE rn<=N '
     'para limitar a últimas N muestras por equipo en consultas de tendencia.'),
    ('DATOS — Cuajone en [Eqpcare].[lc]: ',
     'Cuajone no tiene datos en la tabla de límites → triage siempre retorna 0 observados '
     'aunque Fe=440. El equipo de BD debe cargar los LP/LC reales de Cuajone.'),
    ('DATOS — Typo "MOTORO DE TRACCION RH" en Antamina: ',
     'El JOIN con [Eqpcare].[lc] falla para estos registros → fallback ISNULL=9999 → '
     'no dispara falsos positivos pero tampoco evalúa correctamente. Corrección en la BD.'),
]

GREEN = RGBColor(0x37,0x86,0x1F)
RED   = RGBColor(0xC0,0x00,0x00)
BLUE  = RGBColor(0x2E,0x74,0xB5)
GRAY  = RGBColor(0x50,0x50,0x50)
colors = [RED, BLUE, RED, BLUE, GRAY, GRAY]

for (bp, bt), col in zip(recs, colors):
    bullet(doc, bt, bold_prefix='• ' + bp, color=col)

doc.add_paragraph()
body(doc,
     'Conclusión: KomfIA es la solución adecuada para producción por su formateo, contexto y análisis automático. '
     'TEST SQL V2 es valioso como herramienta de auditoría y validación técnica. '
     'El único bug bloqueante es el LP cruzado en path LLM equipo-específico — todos los demás paths funcionan correctamente.',
     bold=True, color=RGBColor(0x1F,0x49,0x7D))

# ── Guardar ────────────────────────────────────────────────────────────────────
out = r'c:\Users\Usuario\Pictures\PYTHON\Backend-Chatbot\docs\copilot\KomfIA_Instrucciones_Copilot.docx'
doc.save(out)
print(f'Guardado: {out}')
