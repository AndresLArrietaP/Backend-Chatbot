# -*- coding: utf-8 -*-
"""
Genera el documento de Conocimiento "Formatos de Respuesta" para KomfIA Central.
Salida: docs/copilot/knowledge/Formatos_de_Respuesta.docx

POR QUÉ EXISTE: el central tendía a renderizar los datos DOS veces (formato antiguo
horizontal + el esqueleto vertical) y a desviarse de la matriz. Los LLM imitan
EJEMPLOS COMPLETOS mejor que listas de reglas → este doc da la respuesta-modelo de
cada tipo de consulta (✅) y el anti-ejemplo (❌). Es la fuente ESCALABLE del formato:
agregar un tipo de consulta nuevo = agregar su plantilla aquí, sin tocar la instrucción.

Subir como archivo de Conocimiento del agente KomfIA Central.
Ejecutar:  python tools/generar_formatos.py
"""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = pathlib.Path(__file__).resolve().parent.parent / "docs" / "copilot" / "knowledge"


def nuevo_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Cm(1.4)
        s.left_margin = s.right_margin = Cm(1.4)
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(10)
    return d


def H(d, texto, lvl=1):
    p = d.add_paragraph()
    r = p.add_run(texto); r.bold = True
    r.font.size = Pt(13 if lvl == 1 else 11)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if lvl == 1 else RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(8 if lvl == 1 else 5)
    p.paragraph_format.space_after = Pt(2)


def P(d, texto):
    p = d.add_paragraph()
    p.add_run(texto).font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(2)


def CODE(d, texto):
    """Bloque monoespaciado con borde izquierdo — para mostrar la tabla/markdown EXACTA a imitar."""
    for linea in texto.strip("\n").split("\n"):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '6'); left.set(qn('w:color'), 'AAAAAA')
        pBdr.append(left); pPr.append(pBdr)
        r = p.add_run(linea if linea else " ")
        r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


d = nuevo_doc()

p = d.add_paragraph()
r = p.add_run("Formatos de Respuesta — KomfIA (plantillas de presentación)")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
P(d, "Conocimiento de KomfIA Central. Cuando presentes datos, IMITA EXACTAMENTE la plantilla del tipo de "
     "consulta correspondiente. La presentación de datos es UNA SOLA (la matriz vertical): nunca antepongas "
     "ni agregues otra versión. Reutiliza el mismo bloque-matriz para todos los tipos.")

# ── Reglas de oro ────────────────────────────────────────────────────────────
H(d, "Reglas de oro (aplican a TODA respuesta de datos)")
P(d, "1) UNA SOLA presentación EN TABLAS. Puedes abrir con UNA frase de contexto/hallazgo en TEXTO PLANO "
     "(sin tablas, sin datos por muestra), pero la PRIMERA TABLA debe ser «1. Información general»: PROHIBIDA "
     "cualquier tabla ANTES (fechas en filas, «Metales de desgaste», «Estado por metal», «Hallazgos clave»), y "
     "PROHIBIDO repetir los datos en otro formato después. El análisis detallado va DESPUÉS, como viñetas.")
P(d, "2) MATRIZ VERTICAL siempre: los PARÁMETROS van en FILAS y las FECHAS (o la única fecha) en COLUMNAS. "
     "NUNCA pongas las fechas en filas.")
P(d, "3) SEMÁFORO pegado al valor, en la MISMA celda: «5.16 🟥». 🟥 = valor > LC (CRÍTICO) · 🟨 = valor > LP "
     "(PRECAUCIÓN) · 🟩 = OK (en OK puedes dejar la celda limpia, sin emoji). SIN DATO → «—». TBN es inverso "
     "(🟨 si valor < LP). Pb/Sn/TBN no tienen LC.")
P(d, "4) PROHIBIDO: columnas «Estado» separadas, palabras OK/CRÍTICO/PRECAUCIÓN en las celdas de datos, "
     "fusionar los 4 grupos en una sola tabla, y secciones narrativas con tablas (ver anti-ejemplo).")
P(d, "5) Encabezados ≤ 8 caracteres, abreviados con punto (Parámetro→Par., Horómetro→Hor., Horas de "
     "Aceite→Hor. Ace., Fecha→Fec., Estado General→Est. Gen.). «(ppm)» solo en el título de la sección.")
P(d, "6) Fila o columna SIN datos en todas las muestras → se omite, respetando el orden del resto.")

# ── Anti-ejemplo ─────────────────────────────────────────────────────────────
H(d, "❌ ANTI-EJEMPLO — el formato ANTIGUO que NUNCA debes producir")
P(d, "Si te descubres escribiendo cualquiera de estas secciones, DETENTE y bórralas. Son el formato viejo:")
CODE(d, """❌ NUNCA:
«Información General del Equipo» (tabla) +
«Datos Operativos» con # / Fecha / Horómetro / Hrs Aceite / CM / Estado General en FILAS por fecha +
«Metales de Desgaste» con una columna por metal y fechas en filas, o con columnas «Estado X» +
«Índice PQ y Contaminantes», «Aditivos / Salud del Aceite» como tablas separadas horizontales +
«Observaciones Clave de la Tendencia» con 6 puntos numerados.
→ Y LUEGO repetir todo otra vez como matriz. PROHIBIDO. La matriz vertical es la ÚNICA presentación.""")

# ── Último análisis / Condición ──────────────────────────────────────────────
H(d, "✅ ÚLTIMO ANÁLISIS / CONDICIÓN (1 muestra) — respuesta-modelo COMPLETA")
P(d, "Imita esto de principio a fin (datos reales del SQL; la columna de valores se encabeza con la FECHA dd-MMM):")
CODE(d, """A continuación, el último análisis del Motor de Tracción LH del CA3171 (Antapaccay), del 08-Jun.

**1. Información general**
| Campo | Valor |
|---|---|
| Mod./Eqp. | CA3171 |
| Comp. | MT LH |
| Hor. | 35 452 |
| Hor. Ace. | 1 038 |
| CM | M |
| Est. Gen. | 🟩 |
| Fec. | 08-Jun |

**2. Análisis por elemento (ppm)**
| Met. Desg. | LP | LC | 08-Jun |
|---|---|---|---|
| Fe | 200 | 230 | 38.47 |
| PQ | 130 | 150 | 37.6 |
| Cr | 2 | 3 | 0.00 |
| Ni | 2 | 3 | 0.00 |
| Cu | 10 | 15 | 0.00 |
| Pb | 3 | — | 0.20 |
| Sn | 3 | — | 0.00 |
| Al | 2 | 3 | 0.00 |

| Contam. | LP | LC | 08-Jun |
|---|---|---|---|
| Si | 75 | 80 | 24.85 |
| Ca | 18 | 25 | 0.20 |
| Zn | 18 | 25 | 177.43 🟥 |
| K | 2 | 3 | 0.01 |

| Adit. | LP | LC | 08-Jun |
|---|---|---|---|
| B | — | — | 13.12 |
| P | — | — | 404.04 |
| Mg | — | — | 10.67 |

| Salud | LP | LC | 08-Jun |
|---|---|---|---|
| V100 | — | — | 74.43 |
| TBN | — | — | — |

- Estado general OK; todos los metales de desgaste dentro de límite.
- Único punto: Zn 177.43 supera su referencia (aditivo, informativo) → ver recomendación.

🔧 Recomendaciones Técnicas
- Zinc (Zn): <indicio textual del conocimiento «Recomendaciones MT»>.
<párrafo de cierre común, correo una sola vez>

✅ Los equipos consultados se encuentran dentro de límites.""")
P(d, "Notas: si ningún parámetro está fuera de umbral, omite el bloque 🔧 y cierra con ✅. Para CONDICIÓN es "
     "idéntico (1 sola muestra). Las filas Adit./Salud sin LP/LC muestran «—» en esas columnas.")

# ── Tendencia ────────────────────────────────────────────────────────────────
H(d, "✅ TENDENCIA (8 muestras) — respuesta-modelo COMPLETA")
P(d, "⚠️ ES AQUÍ DONDE MÁS SE FALLA. Esto de abajo es TODA la respuesta de tendencia, de principio a fin. "
     "NO escribas NADA antes de «1. Información general» — ni «Tendencia <equipo>», ni «Metales de desgaste y "
     "contaminación», ni «Aditivos e informativos», ni «Observaciones clave de la tendencia», ni ninguna tabla "
     "con fechas en FILAS. La matriz vertical de abajo ES el análisis; las viñetas van al final, no antes.")
P(d, "Una COLUMNA por fecha (dd-MMM), orden cronológico. LP/LC vienen de la vista (Fe_LP, Fe_LC…).")
CODE(d, """A continuación, la tendencia de las 8 últimas muestras del MT LH del CA3171 (Antapaccay).

**1. Información general**
| Campo | Valor |
|---|---|
| Mod./Eqp. | CA3171 |
| Comp. | MT LH |

**2. Tendencia (8 últimas muestras)**
| Par. | 11-May | 14-May | 20-May | 23-May | 26-May | 28-May | 02-Jun | 08-Jun |
|---|---|---|---|---|---|---|---|---|
| Hor. | 34 898 | 34 968 | 35 084 | 35 138 | 35 218 | 35 257 | 35 347 | 35 452 |
| Hor. Ace. | 484 | 554 | 670 | 724 | 804 | 843 | 933 | 1 038 |
| CM | M | M | M | M | M | ADI | M | M |
| Est. Gen. | 🟥 | 🟥 | 🟥 | 🟥 | 🟥 | 🟥 | 🟥 | 🟩 |

**3. Tendencia por elemento (ppm)**
| Met. Desg. | LP | LC | 11-May | 14-May | 20-May | 23-May | 26-May | 28-May | 02-Jun | 08-Jun |
|---|---|---|---|---|---|---|---|---|---|---|
| Fe | 200 | 230 | 80.79 | 95.23 | 110.92 | 74.31 | 93.86 | 100.27 | 77.28 | 38.47 |
| PQ | 130 | 150 | 63.2 | 79.2 | 70.2 | 70.2 | 76.2 | 76.4 | 32.8 | 37.6 |
| Cr | 2 | 3 | 5.16 🟥 | 6.21 🟥 | 7.44 🟥 | 3.57 🟥 | 5.80 🟥 | 6.83 🟥 | 3.95 🟥 | 0.00 |
| Ni | 2 | 3 | 0.14 | 0.00 | 0.29 | 0.77 | 0.48 | 0.50 | 0.00 | 0.00 |
| Cu | 10 | 15 | 0.28 | 0.30 | 0.36 | 0.24 | 0.27 | 0.30 | 0.12 | 0.00 |
| Pb | 3 | — | 0.00 | 0.00 | 0.00 | 4.71 🟨 | 0.10 | 0.08 | 0.32 | 0.20 |
| Sn | 3 | — | 0.00 | 0.00 | 0.00 | 0.00 | 0.00 | 0.00 | 0.00 | 0.00 |
| Al | 2 | 3 | 0.00 | 0.19 | 0.33 | 0.00 | 0.13 | 0.00 | 0.16 | 0.00 |

| Contam. | LP | LC | 11-May | 14-May | 20-May | 23-May | 26-May | 28-May | 02-Jun | 08-Jun |
|---|---|---|---|---|---|---|---|---|---|---|
| Si | 75 | 80 | 23.05 | 18.18 | 20.28 | 23.46 | 23.21 | 20.36 | 22.19 | 24.85 |
| Ca | 18 | 25 | 0.46 | 2.82 | 0.54 | 0.00 | 0.00 | 0.21 | 0.21 | 0.20 |
| Zn | 18 | 25 | 0.00 | 0.12 | 0.00 | 88.74 🟥 | 0.00 | 0.00 | 0.00 | 177.43 🟥 |
| K | 2 | 3 | 0.35 | 0.36 | 0.40 | 0.56 | 0.41 | 0.40 | 0.37 | 0.01 |

(siguen, con el mismo encabezado de fechas, los grupos Adit.: B, P, Mg y Salud: V100, TBN)

```
Cromo (Cr) vs fecha — ppm
 8 |        ×7.44
 7 |  ×6.21      ×6.83
 6 | ×5.16  ×5.80
 5 |              ×3.95
 4 |     ×3.57 ─────────────────── LC 3
 3 |                          ─── LP 2
 2 |
 0 |                    ×0.00
   +11M-14M-20M-23M-26M-28M-02J-08J
```

| Par. | Prom. | σ | Obs. |
|---|---|---|---|
| Fe | 83.9 | 21.3 | dentro de límite, descendente |
| Cr | 4.87 | 2.60 | 7/8 muestras > LC; normaliza 08-Jun |
| Zn | 33.5 | 65.8 | picos anómalos 23-May y 08-Jun |

- Cr persistentemente CRÍTICO 7/8 muestras; solo la última (08-Jun) bajó a 0.
- Zn con picos anómalos (88.74 y 177.43) → posible contaminación/mezcla; merece seguimiento.
- Fe descendente, último 38.47 muy por debajo del LP.

🔧 Recomendaciones Técnicas (sobre la última muestra, 08-Jun)
- <solo si la ÚLTIMA muestra tiene parámetros fuera de umbral; si está en límites, no hay bloque>

📈 Tendencia registrada según muestras disponibles.""")
P(d, "Reglas tendencia: UN solo gráfico ASCII (la métrica más relevante) con líneas LP/LC. La tabla "
     "|Par.|Prom.|σ|Obs.| incluye solo parámetros relevantes. Las recomendaciones evalúan SOLO la última "
     "muestra. Promedio por periodo (mensual/trimestral) solo si lo piden.")

# ── Otros tipos ──────────────────────────────────────────────────────────────
H(d, "TRIAGE / OBSERVADOS de un proyecto")
P(d, "Una fila por equipo-componente observado (no la matriz vertical: aquí el eje es el equipo). Críticos "
     "arriba. Solo columnas de parámetros que alguien supera. Semáforo pegado al valor.")
CODE(d, """De los 54 motores de tracción de Antapaccay, 8 están observados (3 CRÍTICOS, 5 en precaución):

| Equipo | Comp. | Fec. | Fe | Cr | PQ | Pb | Est. Gen. |
|---|---|---|---|---|---|---|---|
| CA3174 | MT RH | 22-May | 311.6 🟥 | 1.8 | 118.4 | — | 🟥 |
| CA3165 | MT RH | 02-Jun | 244.7 🟥 | 2.39 🟨 | 78.0 | — | 🟥 |
| CA3177 | MT LH | 29-May | 146.1 | 1.71 | 201.4 🟥 | — | 🟥 |
| CA3161 | MT LH | 05-Jun | 228.6 🟨 | — | — | — | 🟨 |
| CA3193 | MT RH | 19-May | — | — | — | 3.68 🟨 | 🟨 |

⚠️ Se recomienda seguimiento a estos equipos.""")

H(d, "DIAGNÓSTICO INTEGRAL (un equipo, todos sus componentes)")
P(d, "(a) Resumen ejecutivo: una fila por componente con su Estado y la alerta principal. (b) Por cada "
     "componente OBSERVADO, la matriz vertical de «último análisis» (igual que arriba). Los componentes OK "
     "no necesitan matriz. Indicios formales SOLO para MT; no-MT observado → «monitorear/investigar».")
CODE(d, """Diagnóstico integral del CA3177 (Antapaccay) — 6 compartimientos, fecha 29-May/04-Jun.

| Comp. | Est. Gen. | Alerta principal |
|---|---|---|
| MT LH | 🟥 | PQ 201.4 (LC 150) |
| Sist. Hidráulico | 🟥 | Cu 5.06 (LC 4) |
| Motor | 🟩 | — |
| MT RH | 🟩 | — |
| Rueda Del. LH | 🟩 | — |
| Rueda Del. RH | 🟩 | — |

(luego, solo para MT LH y Sist. Hidráulico, la matriz vertical de su último análisis)
🔧 Recomendaciones: indicio formal de PQ (MT). Hidráulico → monitorear Cu (no-MT, límites por validar).
⚠️ Se recomienda seguimiento a estos equipos.""")

H(d, "BARRIDO de flota (todos los componentes observados de un proyecto)")
P(d, "Agrupa por TIPO de componente; dentro de cada grupo, una fila por equipo observado (estilo triage). "
     "Cierra con un resumen consolidado por tipo. Recuerda: el estado no-MT es preliminar (límites por validar).")
CODE(d, """Barrido de Antapaccay — equipos-componente fuera de límite.

● Motor de Tracción (CRÍTICOS)
| Equipo | Fec. | Métrica | Valor | LP | LC |
|---|---|---|---|---|---|
| CA3165 | 02-Jun | Fe | 244.7 🟥 | 200 | 230 |
| CA3174 | 22-May | Fe | 311.6 🟥 | 200 | 230 |

● Sistema Hidráulico (Cu — límites por validar)
| Equipo | Fec. | Cu | LP | LC |
|---|---|---|---|---|
| CA3196 | 20-May | 21.28 🟥 | 3 | 4 |

| Tipo componente | 🟥 Críticos | 🟨 Prec. | Total |
|---|---|---|---|
| Sistema Hidráulico | 20 | 1 | 21 |
| Motor de Tracción | 3 | 5 | 8 |

⚠️ Se recomienda seguimiento a estos equipos.""")

H(d, "HISTORIAL muestra por muestra")
P(d, "Misma matriz vertical de TENDENCIA (parámetros en filas, fechas en columnas), pero con TODAS las "
     "muestras del compartimiento (incluye DDI; márcalo en la fila CM). Si son muchas, muestra las más "
     "recientes y di cuántas hay en total. NO uses el formato horizontal antiguo.")

P(d, "Para CUALQUIER tipo nuevo de consulta, reutiliza el bloque-matriz vertical de arriba. Si dudas, una "
     "sola tabla por grupo (Met. Desg. / Contam. / Adit. / Salud), parámetros en filas, semáforo en la celda.")

d.save(str(OUT / "Formatos_de_Respuesta.docx"))
print("Guardado:", OUT / "Formatos_de_Respuesta.docx")
