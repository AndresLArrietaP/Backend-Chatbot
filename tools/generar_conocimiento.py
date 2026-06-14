# -*- coding: utf-8 -*-
"""
Genera los documentos de Conocimiento (Knowledge) para Copilot Studio, listos para subir.
Salida en docs/copilot/knowledge/:
  - Recomendaciones_MT.docx          → Conocimientos de KomfIA Central
  - Esquema_y_Patrones_SQL.docx      → Conocimientos de KomfIA SQL
  - Limites_LP_LC_por_Proyecto.docx  → Conocimientos de KomfIA SQL (referencia)

Estructurados para RAG: títulos claros, cada tema en su sección.
Ejecutar:  python tools/generar_conocimiento.py

⚠️ AVISO (2026-06-09): los .docx de knowledge fueron EDITADOS A MANO y validados en producción
(incorporan vw_MuestrasRankeadas / vw_MuestrasEstado, regla DDI asimétrica, agrupación Fe+PQ solo si
ambos observados, nota jun-2026 de límites). Este generador AÚN NO refleja esos cambios → los .docx en
docs/copilot/knowledge/ son la FUENTE CANÓNICA. NO regenerar desde aquí sin antes portar esas ediciones,
o se revertirá contenido validado. Para tocar el conocimiento, edita el .docx directamente.
"""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = pathlib.Path(__file__).resolve().parent.parent / "docs" / "copilot" / "knowledge"
OUT.mkdir(parents=True, exist_ok=True)

AZUL = RGBColor(0x1F, 0x49, 0x7D)


def nuevo_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(11)
    return d


def H(d, t, lvl=1):
    p = d.add_heading(t, level=lvl)
    r = p.runs[0] if p.runs else p.add_run(t)
    r.font.size = Pt({1: 16, 2: 13, 3: 11.5}.get(lvl, 11))
    r.font.color.rgb = AZUL
    return p


def P(d, t, bold=False):
    p = d.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(11)
    return p


def CODE(d, t):
    for ln in t.strip("\n").split("\n"):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Courier New"; r.font.size = Pt(9)
    d.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# DOC 1 — Recomendaciones MT  (Conocimiento de KomfIA Central)
# ══════════════════════════════════════════════════════════════════════════════
# INDICIO por metal = causa + acción específica (lo que VARÍA entre metales).
# El cierre común (acciones generales + contacto) se muestra UNA sola vez al final.
INDICIOS = {
    "Hierro (Fe) y PQ":
        "Alto Hierro y PQ puede indicar un problema en los engranajes o cojinetes. De continuar elevada "
        "la tendencia, solicitar inspección del piñón solar.",
    "Cromo (Cr)":
        "Alto Cromo puede indicar un problema en los rodillos y pistas de rodamientos.",
    "Niquel (Ni)":
        "Alto Níquel puede indicar un problema en los engranajes. De continuar elevada la tendencia, "
        "solicitar inspección del piñón solar.",
    "Cobre (Cu)":
        "Alto Cobre puede indicar un desgaste en las arandelas de empuje (interna y/o externa) o en el "
        "cojinete. Revise la arandela de empuje si encuentra valores altos de Cu/Pb/Sn en conjunto y "
        "reemplácela si hay daño o desgaste excesivo.",
    "Plomo (Pb) y Estaño (Sn)":
        "Alto Plomo y Estaño, acompañado de alto Cu, puede indicar un desgaste en las arandelas de empuje "
        "(interna y/o externa). Revise la arandela de empuje y reemplácela si hay daño o desgaste excesivo.",
    "Silicio (Si)":
        "Alto Silicio probablemente esté asociado a ingreso de contaminación a la caja de engranajes. "
        "Revise la tendencia de Al: si ambas suben en paralelo indica tierra/polvo abrasivo, con "
        "correlación en el incremento de Fe/PQ.",
}
CIERRE_COMUN = (
    "Acortar la frecuencia de monitoreo y programar dializado/cambio de aceite en el próximo PM. "
    "Retirar los 8 tapones magnéticos para inspección y limpieza en busca de particulado anormal. "
    "Para mayor información y detalle, contactar a confiabilidad.operaciones@kmmp.com.pe"
)

d1 = nuevo_doc()
H(d1, "Recomendaciones Técnicas — Motor de Tracción (MT) — KMMP")
P(d1, "Documento de conocimiento para KomfIA. Cada metal tiene su INDICIO específico (causa + acción "
      "particular). El CIERRE COMÚN (acciones generales + contacto) se añade UNA sola vez al final del "
      "bloque, NO por cada metal. Transcribir textualmente.")
H(d1, "Reglas de uso", 2)
P(d1, "• Muestra el indicio de un metal solo cuando su valor supera su LP (PRECAUCIÓN o CRÍTICO).")
P(d1, "• Agrupa metales que comparten indicio: Fe y PQ juntos; Plomo y Estaño juntos.")
P(d1, "• Orden: Fe, PQ, Cromo, Niquel, Cobre, Plomo, Estaño, Silicio.")
P(d1, "• Una viñeta por metal observado con su INDICIO; al final, un párrafo de cierre con el "
      "CIERRE COMÚN (sin etiqueta, solo el párrafo). El correo de contacto aparece UNA sola vez, en ese cierre.")
H(d1, "Indicios por metal (lo que varía)", 2)
for nombre, ind in INDICIOS.items():
    H(d1, f"Indicio — {nombre}", 3)
    P(d1, ind)
H(d1, "Cierre común (párrafo final, añadir UNA sola vez)", 2)
P(d1, CIERRE_COMUN)
d1.save(str(OUT / "Recomendaciones_MT.docx"))
print("Guardado:", OUT / "Recomendaciones_MT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOC 2 — Esquema y Patrones SQL  (Conocimiento de KomfIA SQL)
# ══════════════════════════════════════════════════════════════════════════════
d2 = nuevo_doc()
H(d2, "Esquema de BD y Patrones SQL — bd_kmmp_osconfiabilidad")
P(d2, "Documento de conocimiento para KomfIA SQL. Esquema, relaciones y plantillas de consulta para "
      "generar SQL correcto en SQL Server (Azure). Servidor: serverbd-osconfiabilidad.database.windows.net.")

H(d2, "Tablas y columnas", 2)
H(d2, "[Oil].[LaboratoryData]", 3)
P(d2, "LaboratoryDataId(PK), MiningEquipmentId(FK), CodigoMuestreo, Compartimiento, FechaMuestreo, "
      "Horometro, HorasDeAceite, CM, Grado, Condicion (1=Normal/2=Precaución/3=Crítico — NO usar para estado). "
      "Metales: Fe_ppm, Cu_ppm, Cr_ppm, Al_ppm, Pb_ppm, Sn_ppm, Ni_ppm, Si_ppm, B_ppm, P_ppm. "
      "Físicos: Indice_PQ, V100, TBN, TAN, Agua, Oxidacion.")
H(d2, "[Mine].[MiningEquipment]", 3)
P(d2, "Id(PK, UNIQUEIDENTIFIER), Code (ej: 'CA3160'), MiningProjectId(FK), EquipmentFleetId(FK). "
      "La columna del equipo es ME.[Code]. NUNCA existe EquipmentCode ni EquipmentId.")
H(d2, "[Mine].[MiningProject]", 3)
P(d2, "Id(PK), Name (ej: 'Antapaccay','Antamina','Cerro Verde','Toromocho','Cuajone').")
H(d2, "[Mine].[EquipmentFleet]", 3)
P(d2, "Id(PK), Model (ej: '980E','930E','D475A'), Type, Description.")
H(d2, "[Eqpcare].[lc] — límites LP/LC por proyecto+componente", 3)
P(d2, "Proyecto, COMPONENTE, y por metal: [FIERRO - LP],[FIERRO - LC],[CROMO - LP],[CROMO - LC],"
      "[NIQUEL - LP],[NIQUEL - LC],[COBRE - LP],[COBRE - LC],[SILICIO - LP],[SILICIO - LC],"
      "[ALUMINIO - LP],[ALUMINIO - LC],[PLOMO - LP],[ESTAÑO - LP],[PQ - LP],[PQ - LC],[TBN - LP],[TBN - LC]. "
      "Columnas con espacios SIEMPRE entre corchetes.")
H(d2, "Relaciones (JOINs)", 3)
CODE(d2, "LD.[MiningEquipmentId] = ME.[Id]\nME.[MiningProjectId] = MP.[Id]\nME.[EquipmentFleetId] = EF.[Id]")

H(d2, "VISTAS — usar PRIMERO (resuelven último análisis determinístico + LP/LC + Estado)", 2)
P(d2, "Estas vistas encapsulan el 'último análisis' con desempate determinístico "
      "(FechaMuestreo DESC, LaboratoryDataId DESC), ya excluyen DDI y traen LP/LC + Estado. "
      "Con ellas NO uses ROW_NUMBER. Prefiérelas para condición, triage, ranking y último análisis.")
P(d2, "[dbo].[vw_EstadoActualMT] (Motor de Tracción): columnas Equipo, Proyecto, Modelo, Compartimiento, "
      "FechaMuestreo, métricas, sus LP/LC, Estado_<metal> (Estado_Fe, Estado_Cr, …, Estado_TBN) y "
      "Estado_General (CRITICO/PRECAUCION/OK del equipo).")
P(d2, "[dbo].[vw_UltimoAnalisisAceite] (cualquier componente): última muestra por equipo+compartimiento, "
      "sin Estado. Para 'último análisis' general o componentes que no sean MT.")
H(d2, "Plantillas con vistas", 3)
CODE(d2, """-- Condición de un equipo (MT):
SELECT * FROM [dbo].[vw_EstadoActualMT] WITH (NOLOCK)
WHERE Equipo='CA3171' AND Compartimiento LIKE '%LH';

-- Triage de un proyecto (solo observados); Modelo 980E solo si es Antapaccay:
SELECT TOP 500 * FROM [dbo].[vw_EstadoActualMT] WITH (NOLOCK)
WHERE Proyecto LIKE '%Antapaccay%' AND Modelo LIKE '%980E%' AND Estado_General <> 'OK'
ORDER BY CASE Estado_General WHEN 'CRITICO' THEN 1 ELSE 2 END, Fe_ppm DESC;

-- Ranking top N por metal (TBN: ASC):
SELECT TOP 5 * FROM [dbo].[vw_EstadoActualMT] WITH (NOLOCK)
WHERE Proyecto LIKE '%Antapaccay%' ORDER BY Fe_ppm DESC;""")
P(d2, "TENDENCIA e HISTORIAL NO usan las vistas (necesitan varias muestras): van contra la tabla base "
      "[Oil].[LaboratoryData] con ROW_NUMBER (ver plantillas de abajo). Tendencia excluye DDI; historial los incluye.")

H(d2, "Reglas SQL obligatorias", 2)
P(d2, "• Solo SELECT/CTE. WITH (NOLOCK) en todas las tablas. TOP N (nunca LIMIT). Sin límite: TOP 100. "
      "EXCEPCIÓN — en TRIAGE de un proyecto/flota completa usa TOP 500 (puede haber >100 compartimientos "
      "observados); NUNCA truncar el triage con TOP 100.")
P(d2, "• ROW_NUMBER (solo en tabla base, para tendencia/historial; con vistas NO se usa): "
      "ROW_NUMBER() OVER (PARTITION BY LD.[MiningEquipmentId], LD.[Compartimiento] "
      "ORDER BY LD.[FechaMuestreo] DESC, LD.[LaboratoryDataId] DESC). Las DOS columnas del ORDER BY son "
      "OBLIGATORIAS: sin LaboratoryDataId el 'último' es no-determinista (hay fechas empatadas). "
      "NUNCA TOP dentro de un CTE.")
P(d2, "• Compartimiento con keyword corto, NUNCA frase con 'DE': tracción→'%TRACCION%', "
      "hidráulico→'%HIDRAUL%', rueda→'%RUEDA%', mando final→'%MANDO%', transmisión→'%TRANSMISION%', "
      "motor principal→'%MOTOR%' AND NOT LIKE '%TRACCION%'.")
P(d2, "• Lado: 'izquierdo/izquierda'=LH, 'derecho/derecha'=RH → filtra SOLO las muestras "
      "(AND LD.[Compartimiento] LIKE '%LH'), NUNCA en LimitesLC.")
P(d2, "• Límites LP/LC: SIEMPRE con subconsulta SIN GROUP BY + CROSS JOIN (NUNCA LEFT JOIN exact-match "
      "por COMPONENTE, falla si el naming difiere entre proyectos). El valor real viene del CROSS JOIN; "
      "ISNULL solo es red de seguridad: Fe→200, resto→9999. LP Fe MT: Antapaccay/Antamina=200, "
      "Cerro Verde=80, Toromocho=145.")
P(d2, "• Estado por metal (CASE): >LC=CRÍTICO, >LP=PRECAUCIÓN, resto OK. Con LC: Fe, Cr, Ni, Cu, Si, Al, PQ. "
      "Pb/Sn solo LP. TBN INVERSO: PRECAUCIÓN si valor < [TBN - LP].")
P(d2, "• Excluir post-dializado en triage/tendencia/ranking: "
      "AND (LD.[CM] IS NULL OR LD.[CM] NOT IN ('DDI','DIALIZADO','RELLENO+DIALIZADO')). Historial: incluir todos.")
P(d2, "• Antapaccay explícito → filtrar flota 980E (JOIN EquipmentFleet, EF.[Model] LIKE '%980E%'). "
      "Sin proyecto mencionado, no forzar 980E.")

H(d2, "Plantilla — Condición / Estado de un equipo (último análisis + LP/LC + Estado por metal)", 2)
CODE(d2, """WITH LatestSamples AS (
  SELECT ME.[Code] AS [Equipo], MP.[Name] AS [Proyecto], LD.[Compartimiento], LD.[FechaMuestreo],
    LD.[Horometro], LD.[HorasDeAceite], LD.[CM], LD.[Grado],
    LD.[Fe_ppm],LD.[Cr_ppm],LD.[Ni_ppm],LD.[Cu_ppm],LD.[Pb_ppm],LD.[Sn_ppm],LD.[Si_ppm],
    LD.[Al_ppm],LD.[B_ppm],LD.[P_ppm],LD.[Indice_PQ],LD.[V100],LD.[TBN],
    ROW_NUMBER() OVER (PARTITION BY LD.[MiningEquipmentId],LD.[Compartimiento]
      ORDER BY LD.[FechaMuestreo] DESC) AS rn
  FROM [Oil].[LaboratoryData] LD WITH (NOLOCK)
  INNER JOIN [Mine].[MiningEquipment] ME WITH (NOLOCK) ON LD.[MiningEquipmentId]=ME.[Id]
  INNER JOIN [Mine].[MiningProject] MP WITH (NOLOCK) ON ME.[MiningProjectId]=MP.[Id]
  INNER JOIN [Mine].[EquipmentFleet] EF WITH (NOLOCK) ON ME.[EquipmentFleetId]=EF.[Id]
  WHERE ME.[Code]='CA3171' AND MP.[Name] LIKE '%Antapaccay%' AND EF.[Model] LIKE '%980E%'
    AND LD.[Compartimiento] LIKE '%TRACCION%' AND LD.[Compartimiento] LIKE '%LH'
    AND (LD.[CM] IS NULL OR LD.[CM] NOT IN ('DDI','DIALIZADO','RELLENO+DIALIZADO'))
),
LimitesLC AS (
  SELECT MIN([FIERRO - LP]) AS [FIERRO_LP], MIN([FIERRO - LC]) AS [FIERRO_LC],
         MIN([CROMO - LP]) AS [CROMO_LP], MIN([CROMO - LC]) AS [CROMO_LC],
         MIN([NIQUEL - LP]) AS [NIQUEL_LP], MIN([NIQUEL - LC]) AS [NIQUEL_LC],
         MIN([COBRE - LP]) AS [COBRE_LP], MIN([COBRE - LC]) AS [COBRE_LC],
         MIN([SILICIO - LP]) AS [SILICIO_LP], MIN([SILICIO - LC]) AS [SILICIO_LC],
         MIN([ALUMINIO - LP]) AS [ALUMINIO_LP], MIN([ALUMINIO - LC]) AS [ALUMINIO_LC],
         MIN([PLOMO - LP]) AS [PLOMO_LP], MIN([ESTAÑO - LP]) AS [ESTAÑO_LP],
         MIN([PQ - LP]) AS [PQ_LP], MIN([PQ - LC]) AS [PQ_LC], MAX([TBN - LP]) AS [TBN_LP]
  FROM [Eqpcare].[lc] WITH (NOLOCK)
  WHERE [Proyecto] LIKE '%Antapaccay%' AND [COMPONENTE] LIKE '%TRACCION%'
)
SELECT TOP 1 LS.*, LC.*,
  CASE WHEN LS.[Fe_ppm]>ISNULL(LC.[FIERRO_LC],9999) THEN 'CRÍTICO'
       WHEN LS.[Fe_ppm]>ISNULL(LC.[FIERRO_LP],200)  THEN 'PRECAUCIÓN' ELSE 'OK' END AS [Estado_Fe],
  CASE WHEN LS.[Cr_ppm]>ISNULL(LC.[CROMO_LC],9999) THEN 'CRÍTICO'
       WHEN LS.[Cr_ppm]>ISNULL(LC.[CROMO_LP],9999) THEN 'PRECAUCIÓN' ELSE 'OK' END AS [Estado_Cr]
  -- (repetir CASE por metal: Ni, Cu, Si, Al con LC; Pb/Sn solo LP; PQ con LC; TBN inverso)
FROM LatestSamples LS CROSS JOIN LimitesLC LC WHERE LS.rn=1;""")

H(d2, "Plantilla — Triage (solo observados de un proyecto)", 2)
P(d2, "Igual que Condición, pero SIN equipo fijo y el WHERE final filtra a los que superan algún LP:")
CODE(d2, """FROM LatestSamples LS CROSS JOIN LimitesLC LC
WHERE LS.rn=1 AND (
  LS.[Fe_ppm]>ISNULL(LC.[FIERRO_LP],200) OR LS.[Cr_ppm]>ISNULL(LC.[CROMO_LP],9999) OR
  LS.[Ni_ppm]>ISNULL(LC.[NIQUEL_LP],9999) OR LS.[Cu_ppm]>ISNULL(LC.[COBRE_LP],9999) OR
  LS.[Si_ppm]>ISNULL(LC.[SILICIO_LP],9999) OR LS.[Al_ppm]>ISNULL(LC.[ALUMINIO_LP],9999) OR
  LS.[Pb_ppm]>ISNULL(LC.[PLOMO_LP],9999) OR LS.[Sn_ppm]>ISNULL(LC.[ESTAÑO_LP],9999) OR
  LS.[Indice_PQ]>ISNULL(LC.[PQ_LP],9999) OR
  (LC.[TBN_LP] IS NOT NULL AND LS.[TBN]>0 AND LS.[TBN]<LC.[TBN_LP]) )
ORDER BY LS.[Fe_ppm] DESC;""")

H(d2, "Plantilla — Tendencia (últimas 8 muestras individuales, con LP/LC)", 2)
CODE(d2, """WITH Samples AS (
  SELECT ME.[Code] AS [Equipo], LD.[Compartimiento], LD.[FechaMuestreo],
    LD.[Fe_ppm],LD.[Cr_ppm],LD.[Ni_ppm],LD.[Cu_ppm],LD.[Si_ppm],LD.[Indice_PQ],LD.[TBN],
    ROW_NUMBER() OVER (PARTITION BY LD.[MiningEquipmentId],LD.[Compartimiento]
      ORDER BY LD.[FechaMuestreo] DESC) AS rn
  FROM [Oil].[LaboratoryData] LD WITH (NOLOCK)
  INNER JOIN [Mine].[MiningEquipment] ME WITH (NOLOCK) ON LD.[MiningEquipmentId]=ME.[Id]
  WHERE ME.[Code]='CA3198' AND LD.[Compartimiento] LIKE '%TRACCION%' AND LD.[Compartimiento] LIKE '%RH'
    AND (LD.[CM] IS NULL OR LD.[CM] NOT IN ('DDI','DIALIZADO','RELLENO+DIALIZADO'))
),
LimitesLC AS ( /* igual que arriba, filtrado por el proyecto del equipo */ )
SELECT s.*, LC.[FIERRO_LP],LC.[FIERRO_LC],LC.[CROMO_LP],LC.[CROMO_LC],LC.[PQ_LP],LC.[PQ_LC],LC.[TBN_LP]
FROM Samples s CROSS JOIN LimitesLC LC
WHERE s.rn<=8 ORDER BY s.[Compartimiento], s.[FechaMuestreo] ASC;
-- AVG por periodo SOLO si piden 'mensual/trimestral/anual': EOMONTH/DATEFROMPARTS + GROUP BY.""")

H(d2, "Plantilla — Ranking y Conteo", 2)
CODE(d2, """-- RANKING top N por metal: rn=1 por equipo, ORDER BY metal DESC (TBN ASC), TOP N.
-- CONTEO de flota:
SELECT COUNT(ME.[Id]) AS Total, MP.[Name] AS Proyecto, EF.[Model] AS Modelo
FROM [Mine].[MiningEquipment] ME WITH (NOLOCK)
JOIN [Mine].[EquipmentFleet] EF WITH (NOLOCK) ON EF.[Id]=ME.[EquipmentFleetId]
JOIN [Mine].[MiningProject] MP WITH (NOLOCK) ON MP.[Id]=ME.[MiningProjectId]
WHERE MP.[Name] LIKE '%Antapaccay%' AND EF.[Model] LIKE '%980%'
GROUP BY MP.[Name], EF.[Model];""")

d2.save(str(OUT / "Esquema_y_Patrones_SQL.docx"))
print("Guardado:", OUT / "Esquema_y_Patrones_SQL.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOC 3 — Límites LP/LC por Proyecto  (referencia)
# ══════════════════════════════════════════════════════════════════════════════
d3 = nuevo_doc()
H(d3, "Límites LP/LC por Proyecto — Motor de Tracción (referencia)")
P(d3, "Valores de referencia confirmados en [Eqpcare].[lc] para Motor de Tracción. La consulta SIEMPRE "
      "debe obtener el LP/LC real desde la tabla (CROSS JOIN); esta tabla es solo referencia/validación.")
t = d3.add_table(rows=1, cols=5); t.style = "Table Grid"
hdr = ["Proyecto", "Fe LP", "Fe LC", "Cu LP", "Si LP"]
for i, h in enumerate(hdr):
    c = t.rows[0].cells[i]; c.text = h; c.paragraphs[0].runs[0].bold = True
for fila in [["Antamina","200","233","30","40"], ["Antapaccay","200","230","10","75"],
             ["Cerro Verde","80","168","6","21"], ["Toromocho","145","165","6","21"],
             ["Cuajone","sin datos","—","—","—"]]:
    row = t.add_row()
    for i, v in enumerate(fila):
        row.cells[i].text = v
d3.add_paragraph()
P(d3, "Notas: Cuajone sin datos en lc → 0 observados (correcto, no inventar LP). Cromo MT Antapaccay: "
      "LP=2, LC=3. PQ MT Antapaccay: LP=130, LC=150. Por ahora solo Motor de Tracción; otros "
      "componentes (Hidráulico, Rueda, Motor) se agregarán cuando el área entregue sus límites.")
d3.save(str(OUT / "Limites_LP_LC_por_Proyecto.docx"))
print("Guardado:", OUT / "Limites_LP_LC_por_Proyecto.docx")
