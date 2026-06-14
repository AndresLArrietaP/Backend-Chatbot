# -*- coding: utf-8 -*-
"""
Genera DocumentacionTecnica_BackendChatbot.docx
Ejecutar: python generar_documentacion.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ──────────────────────────────────────────────────────────────────────────────
#  HELPERS  (doc como global — script de generación de un solo uso)
# ──────────────────────────────────────────────────────────────────────────────

def h(titulo, nivel=1):
    """Encabezado de sección."""
    tamanos = {1: 16, 2: 13, 3: 11}
    colores  = {1: (31, 73, 125), 2: (54, 96, 146), 3: (79, 129, 189)}
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.font.name = "Calibri"
    run.font.size = Pt(tamanos.get(nivel, 11))
    run.bold = True
    run.font.color.rgb = RGBColor(*colores.get(nivel, (0, 0, 0)))
    p.paragraph_format.space_before = Pt(14 if nivel == 1 else 9)
    p.paragraph_format.space_after  = Pt(4)


def p(texto, negrita=False, tamaño=11, sangria=0, color=None):
    """Párrafo normal."""
    par = doc.add_paragraph()
    if sangria:
        par.paragraph_format.left_indent = Cm(sangria)
    run = par.add_run(texto)
    run.font.name = "Calibri"
    run.font.size = Pt(tamaño)
    run.bold = negrita
    if color:
        run.font.color.rgb = RGBColor(*color)
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after  = Pt(3)
    return par


def codigo(texto, sangria=0.4):
    """Bloque de código con fondo gris claro y fuente monoespaciada."""
    for linea in texto.strip().splitlines():
        par = doc.add_paragraph()
        par.paragraph_format.left_indent  = Cm(sangria)
        par.paragraph_format.space_before = Pt(1)
        par.paragraph_format.space_after  = Pt(1)
        run = par.add_run(linea if linea.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        pPr = par._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def li(texto, sangria=0.8):
    """Elemento de lista con viñeta simple."""
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(sangria)
    par.paragraph_format.space_before = Pt(1)
    par.paragraph_format.space_after  = Pt(1)
    run = par.add_run(f"•  {texto}")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def tabla(encabezados, filas):
    """Tabla de dos columnas con estilo grilla."""
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    for i, eh in enumerate(encabezados):
        cell = t.rows[0].cells[i]
        cell.text = eh
        run = cell.paragraphs[0].runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.bold = True
    for fila in filas:
        row = t.add_row()
        for i, val in enumerate(fila):
            cell = row.cells[i]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Calibri"
            run.font.size = Pt(10)
    doc.add_paragraph()


def sp():
    """Espacio vertical pequeño."""
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────────────────────
#  PORTADA
# ──────────────────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

pt = doc.add_paragraph()
pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = pt.add_run("Documentación Técnica\nBackend de Análisis de Aceite — KMMP")
rt.font.name = "Calibri"
rt.font.size = Pt(22)
rt.bold = True
rt.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()

ps = doc.add_paragraph()
ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = ps.add_run(
    "API de Consulta en Lenguaje Natural sobre Azure SQL Server\n"
    "Generado: " + datetime.date.today().strftime("%d/%m/%Y")
)
rs.font.name = "Calibri"
rs.font.size = Pt(12)
rs.font.color.rgb = RGBColor(90, 90, 90)

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  1. DESCRIPCIÓN GENERAL
# ──────────────────────────────────────────────────────────────────────────────

h("1. Descripción General")

p(
    "Este backend expone una API REST que recibe preguntas en lenguaje natural "
    "(español) relacionadas con el análisis de aceite y el estado de equipos mineros. "
    "Traduce esas preguntas a SQL mediante un modelo de lenguaje (LLM), ejecuta las "
    "consultas contra Azure SQL Server y devuelve los resultados enriquecidos con "
    "análisis estadístico y sugerencias de visualización."
)

sp()
p("Stack tecnológico principal:", negrita=True)
li("Python 3.11 — FastAPI + Uvicorn")
li("Azure SQL Server — conector pymssql (sin necesidad de ODBC Driver instalado)")
li("LLM primario: Google Gemini 2.5 Flash (hedging paralelo + fallback secuencial)")
li("LLM de respaldo: OpenAI GPT-4o")
li("Desplegado en Render Pro con integración a Microsoft Copilot Studio")

sp()
p("Casos de uso principales:", negrita=True)
li("Triage de observados: detecta equipos con algún parámetro fuera de límite (LP/LC).")
li("Historial crudo: muestra a muestra sin promedios, en orden cronológico.")
li("Tendencia mensual: serie de tiempo con promedios por mes para una métrica.")
li("Comparativas y ranking: ordena equipos por el peor valor de un metal.")
li("Consultas libres en español sobre los esquemas dbo, Oil, Eqpcare, report, Mine.")

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  2. FLUJO DE REQUEST
# ──────────────────────────────────────────────────────────────────────────────

h("2. Arquitectura y Flujo de Request")

p(
    "El endpoint principal POST /human_query ejecuta el siguiente flujo de forma "
    "secuencial. El constraint más importante es el timeout de 240 segundos de "
    "Copilot Studio: cada paso está dimensionado para que el total nunca lo supere."
)

sp()
codigo("""
Usuario (Copilot Studio)
  │
  ▼
POST /human_query  [timeout envolvente: 180s]
  │
  ├─ 1. Recuperar contexto de sesión (GestorContextoConversacional)
  │
  ├─ 2. ¿Puede responderse desde memoria sin SQL nuevo?
  │     ├─ _puede_refinar_desde_memoria()   → filtra/ordena resultado previo en Python
  │     └─ _puede_responder_desde_memoria() → interpreta con LLM, sin nueva query
  │
  ├─ 3. Seleccionar esquema relevante (top 18 tablas por scoring)
  │
  ├─ 4. Generar SQL
  │     ├─ intentar_historial_crudo_directo()  → SQL fijo, sin LLM
  │     ├─ intentar_tendencia_directo()        → SQL fijo, sin LLM
  │     ├─ intentar_triage_directo()           → SQL con límites LP/LC reales
  │     └─ consulta_humana_a_sql()             → LLM (Gemini / OpenAI)
  │
  ├─ 5. Validar y ejecutar SQL  [timeout: 80s]
  │     └─ _blindar_sql() → limpiar → validar tablas → seguridad → ejecutar
  │
  ├─ 6. Reintentos automáticos (máx. 1 en producción, presupuesto: 150s)
  │     ├─ REINTENTO 1: resultado vacío
  │     ├─ REINTENTO 2: agrupación nula (NULL en GROUP BY)
  │     ├─ REINTENTO 3: proyección nula (columnas descriptivas = NULL)
  │     └─ REINTENTO 4: ventana sobre-restringida (IS NOT NULL antes de ROW_NUMBER)
  │
  ├─ 7. Análisis estadístico determinístico (analitica.py)
  │
  ├─ 8. Respuesta en lenguaje natural (LLM o fallback estadístico)
  │
  └─ 9. Guardar turno en sesión (contexto_chat.py)
""")

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  3. ESTRUCTURA DE ARCHIVOS
# ──────────────────────────────────────────────────────────────────────────────

h("3. Estructura de Archivos")

codigo("""
Backend-Chatbot/
├── index.py                        Punto de entrada; arranca uvicorn
├── config.py                       Configuración centralizada desde .env
├── requirements.txt
├── .env                            Variables de entorno (no versionar en producción)
│
├── src/
│   ├── __init__.py                 Fábrica FastAPI; CORS; startup warmup
│   ├── main.py                     Endpoints, retry logic, selección de esquema
│   ├── llm.py                      Heurísticas de dominio; rutas directas de SQL
│   ├── database.py                 Introspección de esquema; validación; ejecución
│   ├── analitica.py                Estadísticas, tendencias, sugerencias de gráfico
│   ├── contexto_chat.py            Memoria conversacional por sesión (TTL + disco)
│   └── providers/
│       ├── factory.py              Selecciona proveedor LLM según LLM_PROVIDER
│       ├── gemini_client.py        Hedge paralelo + fallback secuencial (Gemini)
│       └── openai_client.py        Fallback cuando Gemini no está disponible
│
└── test/
    ├── test_conexionazure.py       Valida conexión a Azure SQL
    ├── test_gemini.py              Prueba el proveedor Gemini
    └── test_openai.py              Prueba el proveedor OpenAI
""")

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  4. MÓDULOS
# ──────────────────────────────────────────────────────────────────────────────

h("4. Módulos — Descripción Detallada")


# ── 4.1 index.py ─────────────────────────────────────────────────────────────
h("4.1  index.py — Punto de entrada", nivel=2)

p(
    "Carga la configuración desde config.py, instancia la aplicación FastAPI "
    "a través de src/__init__.py y arranca uvicorn. Es el único archivo que "
    "define el host y puerto de escucha."
)

codigo("""
# Arranque local
python index.py

# Con hot-reload (desarrollo)
uvicorn index:app --host 127.0.0.1 --port 5000 --reload
""")


# ── 4.2 config.py ────────────────────────────────────────────────────────────
h("4.2  config.py — Configuración centralizada", nivel=2)

p(
    "Clase Config que lee todas las variables de entorno con python-decouple. "
    "Soporta múltiples cadenas de conexión dinámicas (DB_CONN_1, DB_CONN_2, …) "
    "útiles para apuntar a distintos servidores según el cliente activo. "
    "Todos los módulos importan sus valores desde aquí, no directamente del entorno."
)

sp()
p("Variables de entorno clave:", negrita=True)
li("DATABASE_URL: cadena de conexión principal  mssql+pymssql://usuario:clave@servidor/bd")
li("LLM_PROVIDER: gemini | openai")
li("GOOGLE_API_KEY: clave de Google AI Studio (Gemini)")
li("OPENAI_API_KEY: clave de OpenAI (fallback)")
li("TARGET_SCHEMAS: esquemas a introspeccionar, separados por coma  (dbo,Oil,Eqpcare,...)")
li("APP_PROFILE: development | production")


# ── 4.3 src/__init__.py ──────────────────────────────────────────────────────
h("4.3  src/__init__.py — Fábrica de la app y warmup", nivel=2)

p(
    "Crea la instancia FastAPI, registra CORS, monta el router de main.py y "
    "define el evento de startup. Durante el startup se ejecutan en segundo "
    "plano dos tareas de calentamiento críticas para evitar timeouts en el "
    "primer request:"
)

li(
    "database.warmup_connection(): ejecuta SELECT TOP 1 1 para establecer "
    "la conexión TCP/TLS con Azure SQL antes del primer request real. "
    "Sin esto, la primera petición paga ~30s de cold-start."
)
li(
    "database.obtener_esquema_json(): pre-carga el caché de esquema "
    "(~10s en Azure SQL con 80+ tablas) para que el primer request "
    "no pague el costo de introspección."
)

sp()
p(
    "Nota: sin este warmup, el primer request pagaría ~40s de cold-start, "
    "acercándose al límite de 240s de Copilot Studio.",
    color=(140, 80, 0)
)


# ── 4.4 src/main.py ──────────────────────────────────────────────────────────
h("4.4  src/main.py — Router principal", nivel=2)

p(
    "Archivo central de la API. Define todos los endpoints y contiene la lógica "
    "de orquestación: selección de esquema relevante, blindaje de SQL, "
    "reintentos automáticos y construcción de la respuesta final."
)

h("Endpoints expuestos", nivel=3)

tabla(
    ["Método + Ruta", "Descripción"],
    [
        ["GET  /health",                    "Healthcheck básico"],
        ["GET  /llm/ping",                  "Verifica disponibilidad del proveedor LLM"],
        ["GET  /llm/models",                "Lista modelos disponibles"],
        ["GET  /chat/context/{session_id}", "Recupera historial de conversación de una sesión"],
        ["GET  /schema",                    "Devuelve esquema de BD (cacheado en memoria)"],
        ["POST /schema/refresh",            "Fuerza recarga del caché de esquema"],
        ["POST /human_query",               "Endpoint principal: NL → SQL → resultado enriquecido"],
        ["POST /sql",                       "Ejecución directa de SQL (solo lectura, validado)"],
    ]
)

h("Selección de esquema relevante", nivel=3)

p(
    "La base de datos tiene más de 80 tablas. Enviar el esquema completo al LLM "
    "satura el contexto y ralentiza la generación. La función "
    "_seleccionar_esquema_para_llm() selecciona las 18 tablas más relevantes "
    "para la consulta actual usando un sistema de puntuación por tokens:"
)
li("+50 puntos: tabla en la lista de tablas prioritarias de dominio (ej: [Eqpcare].[lc])")
li("+12 puntos: un token de la consulta aparece en el nombre de la tabla")
li("+4 puntos: un token de la consulta aparece en el nombre de alguna columna")
li("+20 puntos: boost semántico de dominio (aceite, análisis, laboratorio)")
li("+25 puntos: boost de subdominio específico (falla, carga útil)")
p(
    "Las tablas puente se inyectan automáticamente si se seleccionó alguna "
    "tabla de análisis que las requiera. Por ejemplo, [Eqpcare].[lc] (tabla "
    "de límites LP/LC) se inyecta siempre que se seleccione [Oil].[LaboratoryData].",
    sangria=0.8
)

h("Reintentos automáticos", nivel=3)

p(
    "Cuando el SQL producido por el LLM genera resultados incorrectos o vacíos, "
    "el sistema puede reescribir y reintentar automáticamente. "
    "En producción, MAX_SQL_RETRIES_TOTAL=1 limita a un solo reintento de "
    "cualquier tipo para no exceder el presupuesto de tiempo (RETRY_TIME_BUDGET=150s)."
)

codigo("""
_retries_total = 0          # contador global por request
_t_inicio_sql  = time.time()

# Condición para cualquier reintento
if _retries_total < MAX_SQL_RETRIES_TOTAL \\
        and (time.time() - _t_inicio_sql) < RETRY_TIME_BUDGET:
    # disparar reintento correspondiente
    _retries_total += 1
""")

sp()
p("Los cuatro tipos de reintento:", negrita=True)
li(
    "REINTENTO 1 – Resultado vacío: el SQL retorna 0 filas. El LLM reescribe "
    "con condiciones más amplias. Excepción crítica: si es un triage de "
    "observados, 0 filas = resultado válido ('ninguno observado') y NO se reintenta."
)
li(
    "REINTENTO 2 – Agrupación nula: GROUP BY con LEFT JOIN que no encontró "
    "coincidencias → aparece una sola fila con NULL en todas las dimensiones. "
    "El LLM reescribe usando INNER JOIN o filtrando NULLs."
)
li(
    "REINTENTO 3 – Proyección nula: columnas descriptivas (nombre, modelo, "
    "proyecto) todas en NULL por un JOIN mal mapeado. El LLM añade COALESCE "
    "o corrige el JOIN."
)
li(
    "REINTENTO 4 – Ventana sobre-restringida: el LLM agrega IS NOT NULL sobre "
    "la métrica antes de calcular ROW_NUMBER, eliminando registros válidos "
    "(ej: NULL es valor legítimo en horómetro). El LLM mueve el filtro después "
    "de la ventana."
)

h("Timeout y cancelación verdadera de threads", nivel=3)

p(
    "La ejecución SQL usa asyncio.wait_for + anyio.to_thread.run_sync con "
    "abandon_on_cancel=True. Esto garantiza que si la consulta tarda más de "
    "DB_QUERY_TIMEOUT segundos, el thread de BD se abandona inmediatamente y "
    "la API devuelve HTTP 408, sin bloquear el event loop de FastAPI."
)

codigo("""
resultado = await asyncio.wait_for(
    anyio.to_thread.run_sync(
        lambda: database.consultar(sql, engine, max_rows=MAX_ROWS_HARD),
        abandon_on_cancel=True  # sin esto el event loop queda bloqueado hasta que BD responde
    ),
    timeout=DB_QUERY_TIMEOUT,  # 80s en Render, 60s en local
)
""")

p(
    "La versión anterior usaba run_in_threadpool de FastAPI (cancellable=False), "
    "lo que impedía que el timeout HTTP saliera antes de que la BD respondiera, "
    "haciendo imposible cumplir con el límite de 240s de Copilot Studio.",
    color=(140, 80, 0)
)


# ── 4.5 src/llm.py ───────────────────────────────────────────────────────────
h("4.5  src/llm.py — Heurísticas y orquestación LLM", nivel=2)

p(
    "Capa intermedia entre la API y el proveedor LLM. Analiza la consulta del "
    "usuario con patrones regex para detectar la intención y enriquecer el prompt "
    "con instrucciones SQL específicas del dominio. Esto compensa las limitaciones "
    "de los modelos con una BD que tiene quirks propios: UUIDs como PK, columnas "
    "con guión (requieren corchetes en TSQL), esquemas separados por área, "
    "valores de compartimiento con nombres compuestos, etc."
)

h("Rutas directas de SQL sin LLM", nivel=3)

p(
    "Para los casos de uso más frecuentes y bien definidos, el módulo genera SQL "
    "directamente en Python sin consultar al LLM. Esto elimina la latencia del "
    "modelo (~20-30s) y los riesgos de alucinación."
)
li("intentar_historial_crudo_directo(): todas las muestras de un equipo en orden cronológico.")
li("intentar_tendencia_directo(): promedio mensual de una métrica (serie de tiempo).")
li("intentar_triage_directo(): qué equipos tienen algún parámetro fuera de límite LP/LC según [Eqpcare].[lc].")

h("Heurísticas de dominio", nivel=3)

tabla(
    ["Heurística", "Cuándo se activa"],
    [
        ["triage_observados",    "\"observados\", \"fuera de límite\", \"estado de los N motores\""],
        ["join_proyecto_modelo", "Aceite + proyecto/modelo mencionados en la misma consulta"],
        ["componentes_modelo",   "\"componentes\" o \"compartimientos\" + equipo/modelo"],
        ["compartimiento_aceite","Tipo de componente + análisis de aceite"],
        ["laboratorio_aceite",   "Muestras, ppm, TBN, laboratorio"],
        ["ultimo_ventana",       "\"último\", \"más reciente\", referencias a última muestra"],
        ["comparativa",          "\"supera\", \"vs\", diferencia entre equipos o periodos"],
        ["continuidad",          "Referencias deícticas (\"de esos resultados\", \"los mismos\")"],
        ["criticidad",           "\"crítico\", \"peor\", \"prioriza\", ordenamiento por severidad"],
        ["tendencia_historica",  "\"tendencia\", \"evolución\", \"mes a mes\", \"histórico\""],
    ]
)

p("Nota sobre compartimientos — valores reales en la BD:", negrita=True, color=(140, 0, 0))
p(
    "Los valores de [Oil].[LaboratoryData].[Compartimiento] tienen nombres compuestos "
    "con preposiciones. Las cláusulas LIKE deben usar un keyword único, no la frase completa:",
    sangria=0.5
)

codigo("""
-- CORRECTO: keyword único que siempre está presente
WHERE Compartimiento LIKE '%TRACCION%'

-- INCORRECTO: el valor real es 'MOTOR DE TRACCION RH' (tiene 'DE' en el medio)
WHERE Compartimiento LIKE '%MOTOR TRACCION%'
""")

p(
    "El mismo criterio aplica para 'HIDRAUL%', 'RUEDA%', etc. "
    "Las heurísticas en llm.py inyectan el LIKE correcto en el prompt.",
    sangria=0.5
)


# ── 4.6 src/database.py ──────────────────────────────────────────────────────
h("4.6  src/database.py — Acceso a datos y validación SQL", nivel=2)

p(
    "Gestiona la conexión a la base de datos, la introspección del esquema, "
    "la validación de SQL y la ejecución de consultas. "
    "La BD es de solo lectura: no hay permisos DDL ni DML."
)

h("Validación de seguridad SQL (_blindar_sql)", nivel=3)

p("La función aplica las siguientes comprobaciones en orden, antes de ejecutar:")
li("limpiar_sql(): normaliza espacios, saltos de línea y caracteres Unicode.")
li("sanear_explain(): elimina prefijos EXPLAIN / EXPLAIN ANALYZE.")
li("calificar_tablas(): valida que todas las tablas estén en el esquema permitido.")
li("hacer_left_join_nullsafe_mssql(): convierte LEFT JOINs vacíos en INNER JOIN cuando corresponde.")
li("hacer_groupby_nullsafe_mssql(): añade COALESCE en columnas descriptivas del GROUP BY.")
li("es_select_seguro(): rechaza INSERT, UPDATE, DELETE, DROP, EXEC, xp_ (extensiones del sistema).")
li("restringir_a_tablas(): verifica que el SQL no acceda a tablas no autorizadas.")

h("Caché de esquema", nivel=3)

p(
    "La introspección del esquema cuesta ~10s en Azure SQL (80+ tablas y columnas). "
    "obtener_esquema_json() cachea el resultado en memoria con clave compuesta por "
    "los esquemas solicitados. Se invalida con POST /schema/refresh. "
    "El startup warmup en src/__init__.py pre-carga este caché antes del primer request."
)

h("Exclusión de CTEs en validación de tablas", nivel=3)

p(
    "Las consultas WITH ... AS (CTE) declaran nombres virtuales que aparecen en "
    "FROM/JOIN pero no son tablas reales de la BD. La función "
    "_recopilar_nombres_cte() extrae estos nombres para excluirlos del allowlist "
    "y evitar falsos rechazos del validador."
)

codigo("""
-- Ejemplo: 'LatestSamples' y 'LimitesLC' son CTEs, no tablas reales.
-- Sin esta exclusión, el validador las rechazaría por no estar en el esquema.
WITH LatestSamples AS (
    SELECT ... FROM [Oil].[LaboratoryData]
),
LimitesLC AS (
    SELECT ... FROM [Eqpcare].[lc]
)
SELECT * FROM LatestSamples LS
LEFT JOIN LimitesLC LC ON ...
""")


# ── 4.7 src/providers/gemini_client.py ───────────────────────────────────────
h("4.7  src/providers/gemini_client.py — Cliente Gemini", nivel=2)

p(
    "Cliente principal del LLM. Implementa una estrategia de dos fases para "
    "minimizar la latencia manteniendo resiliencia ante fallos individuales de modelos."
)

h("Fase 1 — Hedge paralelo", nivel=3)

p(
    "Lanza peticiones en paralelo a varios modelos Gemini (ej: 2.5-flash y 2-flash). "
    "El primero en responder con SQL válido cancela el resto. "
    "Esto reduce la latencia P50 sin aumentar el costo P95."
)

codigo("""
# Lanzar hedge paralelo
tareas = [asyncio.create_task(_llamar_modelo(m, prompt)) for m in modelos_hedge]

# Esperar al primero que responda correctamente
ganador = await _primera_respuesta_valida(tareas)
# Las demás tareas se cancelan al salir
""")

h("Fase 2 — Fallback secuencial", nivel=3)

p(
    "Si el hedge falla (todos los modelos devuelven error o SQL inválido), "
    "se intenta secuencialmente el resto de la lista con timeout extendido (×1.4). "
    "Los modelos que devolvieron HTTP 503 se omiten para no perder tiempo."
)

codigo("""
# modelos_503 es un set por-request, no global.
# Si un modelo se recupera entre requests, vuelve a estar disponible.
modelos_503: set[str] = set()

for modelo in modelos_fallback:
    if modelo in modelos_503:
        continue  # saltear modelos saturados en este request
    try:
        resultado = await _llamar_modelo(modelo, prompt, timeout=timeout * 1.4)
        if resultado_valido(resultado):
            return resultado
    except Error503:
        modelos_503.add(modelo)
""")

h("Detección de SQL truncado", nivel=3)

p(
    "La función _sql_parece_incompleto() detecta si el LLM truncó la respuesta "
    "antes de completar el SQL. Señales verificadas:"
)
li("Palabras clave SQL al final de la última línea sin completar (SELECT, FROM, WHERE solos).")
li("Paréntesis o corchetes sin cerrar.")
li("Cadenas de texto sin cerrar (número impar de comillas).")
li("Alias vacíos (AS al final sin nombre de alias).")


# ── 4.8 src/providers/openai_client.py ───────────────────────────────────────
h("4.8  src/providers/openai_client.py — Cliente OpenAI", nivel=2)

p(
    "Cliente de respaldo para cuando Gemini no está disponible o falla de forma "
    "persistente. Implementa la misma interfaz que gemini_client.py, por lo que "
    "es intercambiable a través de factory.py sin cambios en el resto del código. "
    "Se activa configurando LLM_PROVIDER=openai en el archivo .env."
)


# ── 4.9 src/providers/factory.py ─────────────────────────────────────────────
h("4.9  src/providers/factory.py — Fábrica de proveedor LLM", nivel=2)

p(
    "Lee la variable LLM_PROVIDER del entorno y devuelve la instancia del cliente "
    "correspondiente. Se instancia una sola vez al importar src/llm.py, "
    "evitando múltiples inicializaciones."
)

codigo("""
# En src/llm.py (al importar el módulo)
_proveedor = _obtener_proveedor()  # GeminiProvider u OpenAIProvider
""")


# ── 4.10 src/contexto_chat.py ─────────────────────────────────────────────────
h("4.10  src/contexto_chat.py — Memoria conversacional", nivel=2)

p(
    "Gestiona la memoria de conversación por sesión. Permite al chatbot mantener "
    "contexto entre preguntas sucesivas del mismo usuario, resolviendo referencias "
    "como 'de esos resultados, quédate solo con los críticos' o "
    "'¿y el equipo anterior cómo está?'."
)

h("Configuración", nivel=3)

tabla(
    ["Variable de entorno", "Descripción"],
    [
        ["CONTEXTO_CHAT_TTL_MINUTOS",         "Vida de una sesión inactiva (default: 45 min)"],
        ["CONTEXTO_CHAT_MAX_TURNOS",          "Máximo de turnos por sesión, ventana FIFO (default: 8)"],
        ["CONTEXTO_CHAT_MAX_CARACTERES",      "Máximo de caracteres inyectados al LLM (default: 5 000)"],
        ["CONTEXTO_CHAT_MAX_FILAS_RESULTADO", "Máximo de filas guardadas por turno (default: 200)"],
        ["CONTEXTO_CHAT_PERSISTIR_ARCHIVO",   "true → persiste a .cache/contexto_chat.json"],
    ]
)

h("Clases principales", nivel=3)

li("TurnoConversacion: snapshot de un turno (consulta del usuario, SQL generado, respuesta textual, filas de resultado).")
li("ConversacionActiva: colección de turnos de una sesión con timestamp de última actividad para el cálculo de TTL.")
li(
    "GestorContextoConversacional: gestor principal. Thread-safe con threading.Lock "
    "(no asyncio.Lock, porque es invocado tanto desde coroutines FastAPI como desde "
    "threads síncronos de SQLAlchemy)."
)

h("Persistencia atómica a disco", nivel=3)

p(
    "Para evitar corrupción del JSON si el proceso muere durante la escritura, "
    "se usa una escritura atómica: se escribe a un archivo .tmp y luego se "
    "renombra al definitivo. Si el proceso falla, el archivo anterior queda intacto."
)

codigo("""
tmp = ruta_archivo.with_suffix(".tmp")
tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
tmp.replace(ruta_archivo)  # operación atómica en la mayoría de sistemas de archivos
""")


# ── 4.11 src/analitica.py ─────────────────────────────────────────────────────
h("4.11  src/analitica.py — Análisis estadístico", nivel=2)

p(
    "Procesa el resultado SQL de forma determinística (sin LLM) para generar "
    "estadísticas descriptivas, detectar tendencias y sugerir el tipo de gráfico "
    "más apropiado para los datos devueltos."
)

h("Funciones principales", nivel=3)

li(
    "generar_analisis_resultado(): calcula media, mediana, desviación estándar, "
    "mínimo y máximo por columna numérica. Detecta tendencia creciente, "
    "decreciente o estable comparando primera y segunda mitad de los datos."
)
li(
    "renderizar_resumen_analitico(): construye un texto conciso con los hallazgos "
    "para usarlo como contexto en la respuesta del LLM."
)
li(
    "generar_chart_url(): genera URL de QuickChart.io para visualizar los datos. "
    "Estrategia: POST primero (URL permanente corta), GET como fallback "
    "para datasets pequeños (la URL GET se vuelve demasiado larga con >20 filas)."
)

h("Selección de columnas numéricas", nivel=3)

p(
    "La función _buscar_columnas_numericas() aplica un threshold del 45%: "
    "una columna se considera numérica si al menos el 45% de sus valores no son NULL. "
    "Además aplica scoring de dominio para priorizar métricas de análisis de aceite:"
)
li("+20 puntos: columna contiene 'ppm', 'tbn', 'vis', 'tan' (métricas de aceite)")
li("-20 puntos: columna contiene 'id', 'code', 'codigo' (identificadores, no métricas)")

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  5. API ENDPOINTS — DETALLE
# ──────────────────────────────────────────────────────────────────────────────

h("5. Detalle del Endpoint Principal — POST /human_query")

h("Request", nivel=2)

codigo("""
POST /human_query
Content-Type: application/json
x-api-key: <API_KEY>           (opcional si API_KEY no está configurada)

{
  "pregunta":               "¿Cuántos motores de tracción de Antapaccay están observados?",
  "session_id":             "usuario-copilot-abc123",
  "max_rows":               100,
  "incluir_respuesta_texto": false,
  "db_key":                 null   (opcional: seleccionar conexión alternativa)
}
""")

h("Response", nivel=2)

codigo("""
HTTP 200 OK

{
  "sql":        "WITH Observados AS (...) SELECT Equipo, Fe_ppm, ... FROM ...",
  "rows":       [ {"Equipo": "CAT 793F-01", "Fe_ppm": 215, ...}, ... ],
  "row_count":  7,
  "analisis": {
    "columnas": {
      "Fe_ppm": { "media": 198.4, "max": 312, "min": 142, "tendencia": "creciente" }
    }
  },
  "respuesta":  null,
  "session_id": "usuario-copilot-abc123",
  "executed":   true,
  "chart_url":  null
}
""")

h("Códigos de error", nivel=2)

tabla(
    ["Código HTTP", "Significado"],
    [
        ["400", "SQL generado es inválido o rechazado por seguridad"],
        ["408", "Timeout de ejecución SQL (DB_QUERY_TIMEOUT superado)"],
        ["429", "Timeout total del request (REQUEST_TIMEOUT superado)"],
        ["500", "Error interno no manejado"],
        ["503", "Todos los modelos LLM devolvieron error"],
    ]
)

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  6. VARIABLES DE ENTORNO
# ──────────────────────────────────────────────────────────────────────────────

h("6. Variables de Entorno Críticas")

p(
    "Los valores indicados son los configurados en Render (producción). "
    "Están optimizados para no superar el timeout de 240s de Copilot Studio."
)

tabla(
    ["Variable", "Valor producción", "Por qué"],
    [
        ["DB_QUERY_TIMEOUT",       "80",    "20s causaba timeout en queries de 40-60s"],
        ["SQL_MAX_OUTPUT_TOKENS",  "8000",  "3400 truncaba CTEs con 4+ JOINs"],
        ["SQL_REPAIR_MAX_TOKENS",  "4096",  "Prompts de reparación para CTEs complejas"],
        ["REQUEST_TIMEOUT",        "180",   "Responde antes del corte de 240s de Copilot"],
        ["RETRY_TIME_BUDGET",      "150",   "Si ya se gastaron 150s en SQL, no se reintenta"],
        ["MAX_SQL_RETRIES_TOTAL",  "1",     "Máximo un reintento de cualquier tipo"],
        ["GENERAR_RESPUESTA_TEXTO","false", "Copilot tiene su propio LLM, segunda llamada (~30s) es redundante"],
        ["LLM_PER_MODEL_TIMEOUT",  "45.0",  "Flash tarda ~15-30s en prompts de triage"],
        ["LLM_TOTAL_TIMEOUT",      "90.0",  "Hedge + fallback secuencial"],
        ["DB_DIALECT",             "mssql", "SQL Server usa TOP N, no LIMIT"],
        ["TARGET_SCHEMAS",         "dbo,Oil,Eqpcare,report,general,module,Mine,Invertex", "Esquemas de la BD"],
        ["GEMINI_MODEL_CANDIDATES","gemini-2.5-flash,gemini-2-flash,gemini-2-flash-lite", "Pro removido (503 consistente)"],
    ]
)

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  7. SEGURIDAD SQL
# ──────────────────────────────────────────────────────────────────────────────

h("7. Modelo de Seguridad SQL")

p(
    "El SQL es generado por un LLM a partir de texto libre del usuario. "
    "La capa de seguridad es crítica y se aplica en múltiples niveles:"
)

li("Solo se permiten sentencias SELECT y CTEs (WITH ... AS ...).")
li("Se rechaza cualquier sentencia DDL (CREATE, ALTER, DROP) o DML (INSERT, UPDATE, DELETE).")
li("Se rechaza EXEC, xp_, procedimientos almacenados y extensiones del sistema.")
li("Todas las tablas referenciadas se validan contra un allowlist por esquema.")
li("El número de filas retornadas tiene un límite duro configurable (MAX_ROWS_HARD, default 2000).")
li(
    "Los literales de cadena se enmascaran antes del scanner de seguridad, "
    "evitando falsos positivos como SELECT * WHERE descripcion='DROP TABLE test'."
)
li(
    "La API puede protegerse con API_KEY en el header x-api-key. "
    "Si API_KEY no está configurada en el entorno, la API acepta peticiones sin clave "
    "(modo desarrollo o entorno cerrado de Copilot)."
)

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  8. DEPENDENCIAS
# ──────────────────────────────────────────────────────────────────────────────

h("8. Dependencias Principales")

tabla(
    ["Paquete", "Uso"],
    [
        ["fastapi",                "Framework web y definición de endpoints"],
        ["uvicorn",                "Servidor ASGI"],
        ["sqlalchemy",             "Introspección de esquema y ejecución de queries"],
        ["pymssql",                "Conector Azure SQL Server (sin ODBC Driver)"],
        ["google-generativeai",    "Cliente Gemini"],
        ["openai",                 "Cliente OpenAI (fallback)"],
        ["python-decouple",        "Lectura de variables de entorno"],
        ["anyio",                  "Threads con abandon_on_cancel para timeout verdadero"],
        ["pydantic",               "Validación de modelos de request/response"],
    ]
)

p(
    "Nota sobre pymssql: se usa en lugar de pyodbc porque Render (Linux) no tiene "
    "permisos para instalar el ODBC Driver de Microsoft. pymssql se conecta "
    "directamente vía protocolo TDS sin dependencias del sistema operativo.",
    color=(140, 80, 0)
)

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  9. EJECUCIÓN Y PRUEBAS
# ──────────────────────────────────────────────────────────────────────────────

h("9. Ejecución Local y Pruebas")

h("Instalación", nivel=2)
codigo("""
pip install -r requirements.txt
""")

h("Arranque del servidor", nivel=2)
codigo("""
# Modo desarrollo (hot-reload)
uvicorn index:app --host 127.0.0.1 --port 5000 --reload

# O directamente
python index.py
""")

h("Scripts de prueba", nivel=2)
codigo("""
# Validar conexión a Azure SQL
python test/test_conexionazure.py

# Probar proveedor Gemini
python test/test_gemini.py

# Probar proveedor OpenAI
python test/test_openai.py
""")

h("Probar el endpoint manualmente", nivel=2)
codigo("""
curl -X POST http://127.0.0.1:5000/human_query \\
  -H "Content-Type: application/json" \\
  -d '{
    "pregunta": "Dame los 5 motores con más hierro en Antapaccay",
    "session_id": "test-local",
    "max_rows": 10,
    "incluir_respuesta_texto": false
  }'
""")

doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
#  10. NOTAS PARA EL EQUIPO DE TI
# ──────────────────────────────────────────────────────────────────────────────

h("10. Notas para el Equipo de TI")

h("Base de datos", nivel=2)
li("El usuario de BD tiene permisos de solo lectura (SELECT). No requiere ni tiene permisos DDL/DML.")
li("El conector es pymssql en el puerto TCP 1433. No requiere instalación de ODBC Driver en el servidor.")
li("La cadena de conexión va en la variable DATABASE_URL del archivo .env, nunca hardcodeada en el código.")

h("Despliegue en Render", nivel=2)
li("Instancia Render Pro (512 MB RAM, 0.5 vCPU). Single worker de uvicorn.")
li("Las variables de entorno se gestionan en el panel de Render, no en el repositorio.")
li("El archivo .env del repositorio contiene solo valores para desarrollo local.")
li("Los logs de arranque correctos son: '[DB] warmup completado' y '[startup] Caché de esquema pre-cargado.'")

h("Integración con Copilot Studio", nivel=2)
li("Copilot Studio llama a POST /human_query como herramienta HTTP personalizada.")
li(
    "Configurar incluir_respuesta_texto: false en la llamada desde Copilot. "
    "Esto evita una segunda llamada al LLM para síntesis (~30s extra), "
    "ya que Copilot tiene su propio sistema de respuesta."
)
li("El timeout de Copilot Studio es de 240s. El backend responde en menos de 180s por diseño.")
li("Si el backend tarda más de 180s, devuelve HTTP 429 con mensaje estructurado antes del corte de Copilot.")

h("Claves de API", nivel=2)
li("GOOGLE_API_KEY: clave de Google AI Studio. Renovar si aparecen errores 403 persistentes.")
li("OPENAI_API_KEY: solo necesaria si LLM_PROVIDER=openai.")
li("API_KEY: clave interna del backend. Si no está en el entorno, la API es pública (aceptable en entorno cerrado de Copilot).")

h("Actualización de modelos Gemini", nivel=2)
p(
    "Si Google depreca algún modelo, actualizar GEMINI_MODEL_CANDIDATES en el panel de Render. "
    "La variable acepta una lista separada por comas. "
    "No se requiere redespliegue; basta con reiniciar el servicio tras guardar la variable."
)


# ──────────────────────────────────────────────────────────────────────────────
#  GUARDAR
# ──────────────────────────────────────────────────────────────────────────────

ruta = r"c:\Users\Usuario\Pictures\PYTHON\Backend-Chatbot\docs\arquitectura\DocumentacionTecnica_BackendChatbot.docx"
doc.save(ruta)
print(f"[OK] Documento generado: {ruta}")
