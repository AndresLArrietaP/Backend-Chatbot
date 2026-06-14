# -*- coding: utf-8 -*-
"""
Genera docs/arquitectura/PLAN_PIVOTE_MULTIAGENTE.docx a partir del .md fuente.

El .md es la fuente única de verdad; este script lo renderiza a Word con formato
corporativo (mismo estilo que generar_instrucciones.py). Conversor markdown ligero:
soporta encabezados (#/##/###), listas (-, •, numeradas), bloques de código (```),
tablas (| ... |), citas (>), reglas (---), negrita (**), e inline code (`...`).

Ejecutar:  python tools/generar_plan_pivote.py
"""
import re
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = pathlib.Path(__file__).resolve().parent.parent
MD_IN = BASE / "docs" / "arquitectura" / "PLAN_PIVOTE_MULTIAGENTE.md"
DOCX_OUT = BASE / "docs" / "arquitectura" / "PLAN_PIVOTE_MULTIAGENTE.docx"

AZUL1 = RGBColor(0x1F, 0x49, 0x7D)
AZUL2 = RGBColor(0x2E, 0x74, 0xB5)
AZUL3 = RGBColor(0x24, 0x63, 0x9E)
GRIS = RGBColor(0x40, 0x40, 0x40)
NEGRO = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.2)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


# ── Helpers de formato ──────────────────────────────────────────────────────────
def heading(text, level):
    p = doc.add_heading(text, level=min(level, 3))
    run = p.runs[0] if p.runs else p.add_run(text)
    size, color = {1: (16, AZUL1), 2: (13, AZUL2), 3: (11.5, AZUL3)}.get(level, (11, AZUL3))
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)


def _runs_con_inline(par, texto):
    """Renderiza **negrita** e `inline code` dentro de un párrafo."""
    # Tokeniza por **...** y `...`
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", texto):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Courier New"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)
        else:
            par.add_run(tok)


def parrafo(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    _runs_con_inline(p, texto)


def vineta(texto, nivel=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 + nivel * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.add_run("• ").bold = True
    _runs_con_inline(p, texto)


def numerada(num, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.add_run(f"{num}. ").bold = True
    _runs_con_inline(p, texto)


def bloque_codigo(lineas):
    for ln in lineas:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Courier New"; r.font.size = Pt(8.5); r.font.color.rgb = NEGRO
    doc.add_paragraph()


def regla():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '1'); b.set(qn('w:color'), '2E74B5')
    pBdr.append(b); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)


def cita(texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '8'); left.set(qn('w:color'), '2E74B5')
    pBdr.append(left); pPr.append(pBdr)
    r = p.add_run("")
    _runs_con_inline(p, texto)
    for rr in p.runs:
        rr.italic = True


def tabla(filas):
    """filas = lista de listas de celdas (strings). Primera fila = encabezado."""
    if not filas:
        return
    ncols = max(len(f) for f in filas)
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    for i, fila in enumerate(filas):
        celdas = t.add_row().cells
        for j in range(ncols):
            val = fila[j] if j < len(fila) else ""
            celdas[j].text = ""
            par = celdas[j].paragraphs[0]
            _runs_con_inline(par, val)
            for rr in par.runs:
                rr.font.size = Pt(9.5)
                if i == 0:
                    rr.bold = True
    doc.add_paragraph()


# ── Parser de markdown ──────────────────────────────────────────────────────────
def parse_tabla(bloque):
    filas = []
    for ln in bloque:
        ln = ln.strip().strip("|")
        if re.match(r"^[\s:\-|]+$", ln):  # separador ---|---
            continue
        filas.append([c.strip() for c in ln.split("|")])
    return filas


def render(md_text):
    lineas = md_text.split("\n")
    i = 0
    tabla_buf = []

    def flush_tabla():
        nonlocal tabla_buf
        if tabla_buf:
            tabla(parse_tabla(tabla_buf))
            tabla_buf = []

    while i < len(lineas):
        ln = lineas[i]
        s = ln.strip()

        # Tabla (acumular líneas que empiezan con |)
        if s.startswith("|"):
            tabla_buf.append(ln)
            i += 1
            continue
        else:
            flush_tabla()

        # Bloque de código
        if s.startswith("```"):
            code = []
            i += 1
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                code.append(lineas[i])
                i += 1
            bloque_codigo(code)
            i += 1
            continue

        # Regla horizontal
        if s in ("---", "***", "___"):
            regla()
            i += 1
            continue

        # Encabezados
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            heading(m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        # Cita
        if s.startswith(">"):
            cita(s.lstrip(">").strip())
            i += 1
            continue

        # Lista numerada
        m = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m:
            numerada(m.group(1), m.group(2).strip())
            i += 1
            continue

        # Viñeta (con nivel por indentación)
        m = re.match(r"^(\s*)[-•]\s+(.*)$", ln)
        if m:
            nivel = len(m.group(1)) // 2
            vineta(m.group(2).strip(), nivel)
            i += 1
            continue

        # Línea vacía
        if not s:
            i += 1
            continue

        # Párrafo normal
        parrafo(s)
        i += 1

    flush_tabla()


# ── Portada ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KomfIA — Plan de Evolución"); r.font.size = Pt(24); r.bold = True; r.font.color.rgb = AZUL1
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pivote a Arquitectura Multiagente (Opus genera SQL)"); r.font.size = Pt(14); r.font.color.rgb = AZUL2
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Komatsu Mitsui Maquinarias Perú S.A. — Área OIS"); r.font.size = Pt(11); r.font.color.rgb = GRIS
doc.add_paragraph()
regla()
doc.add_page_break()

# ── Render del markdown ───────────────────────────────────────────────────────────
md = MD_IN.read_text(encoding="utf-8")
# Quitar el frontmatter de título duplicado del .md (primera línea "# ...") si existe,
# ya que la portada ya lo cubre. Mantener el resto.
render(md)

doc.save(str(DOCX_OUT))
print(f"Guardado: {DOCX_OUT}")
